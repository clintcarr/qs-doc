from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
import os
import get_qlik_sense
import shutil
import sys
import argparse
# import server_details
import create_QSXls
import datetime


from clint.textui import progress

year = datetime.date.today().year
month = datetime.date.today().month
day = datetime.date.today().day

parser = argparse.ArgumentParser()
parser.add_argument('--server', help='Qlik Sense Server to connect to.')
parser.add_argument('--certs', help='Path to certificates.')
parser.add_argument('--user', help='Username in format domain\\user')
parser.add_argument('--wmi', help='True/False, set to True to collect WMI information from servers. This switch requires windows authentication (username and password)')
parser.add_argument('--password', help='Password of user.')
args = parser.parse_args()

try:
    os.remove("Qlik Sense Site.docx")
except FileNotFoundError:
    pass

shutil.copy('./word_template/Qlik Sense Site.docx', 'QSDoc_{0}_{1}_{2}.docx'.format(year, month, day))

document = Document('QSDoc_{0}_{1}_{2}.docx'.format(year, month, day))
sections = document.sections
section = sections[0]

def connect():
    """
    Attempt to connect to the Qlik Sense server specified in the arguments.
    """
    # print ('Testing connection..')
    get_qlik_sense.get_about()

def servernode():
    """
    Generate the Node detail information and write to the output (word document)
    """
    document.add_heading('Qlik Sense Site', 0)
    document.add_heading('Site details', 1)
    servernodes = get_qlik_sense.get_servernode()
    num_of_nodes = len(servernodes)
    table = document.add_table(rows=num_of_nodes+1, cols=4)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Node name'
    row.cells[1].text = 'Host name'
    row.cells[2].text = 'Service Cluster'
    row.cells[3].text = 'Failover Candidate'

    for node in range(num_of_nodes):
        row = table.rows[node+1]
        row.cells[0].text = str(servernodes[node][0])
        row.cells[1].text = str(servernodes[node][1])
        row.cells[2].text = str(servernodes[node][2])

def engine():
    """
    Collect and collate the Qlik Sense Engine information
    """
    document.add_heading('Engine details', 1)

    engine_metrics = ['customProperties','listenerPorts','autosaveInterval', 'tableFilesDirectory', 'genericUndoBufferMaxSize',  'documentTimeout','documentDirectory',
                      'allowDataLineage', 'qrsHttpNotificationPort', 'standardReload',
                      'workingSetSizeLoPct', 'workingSetSizeHiPct', 'workingSetSizeMode','cpuThrottlePercentage', 'maxCoreMaskPersisted', 'maxCoreMask',
                      'maxCoreMaskHiPersisted', 'maxCoreMaskHi','objectTimeLimitSec', 'exportTimeLimitSec', 'reloadTimeLimitSec',
                      'hyperCubeMemoryLimit', 'exportMemoryLimit', 'reloadMemoryLimit', 'createSearchIndexOnReloadEnabled', 'hostname',
                      'globalLogMinuteInterval','auditActivityLogVerbosity','auditSecurityLogVerbosity','serviceLogVerbosity','systemLogVerbosity','performanceLogVerbosity',
                      'qixPerformanceLogVerbosity','auditLogVerbosity','sessionLogVerbosity','trafficLogVerbosity']

    enginenodes = get_qlik_sense.get_engine()
    num_of_engines = len(enginenodes)
    num_of_engine_metrics = len(engine_metrics)
    table = document.add_table(rows=num_of_engine_metrics+1, cols=num_of_engines+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_engines):
        row.cells[item+1].text = enginenodes[item][36]
    for item in range(num_of_engine_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(engine_metrics[item])
        for enginenode in range(num_of_engines):
            row.cells[enginenode+1].text = str(enginenodes[enginenode][item])

    document.add_page_break()

def proxy():
    """
    Collect and collate the Qlik Sense Proxy information
    """
    document.add_heading('Proxy details', 1)
    proxy_metrics = ['customProperties','listenPort','restListenPort', 'allowHttp','unencryptedListenPort','authenticationListenPort',
                     'kerberosAuthentication', 'unencryptedAuthenticationListenPort', 'keepAliveTimeoutSeconds', 'maxHeaderSizeBytes',
                     'maxHeaderLines', 'hostName', 'logVerbosityAuditActivity', 'logVerbosityAuditSecurity', 'logVerbosityService',
                     'logVerbosityAudit', 'logVerbosityPerformance', 'logVerbositySecurity', 'logVerbositySystem','performanceLoggingInterval']

    proxynodes = get_qlik_sense.get_proxy()
    num_of_proxy = len(proxynodes)
    num_of_proxy_metrics = len(proxy_metrics)
    table = document.add_table(rows=num_of_proxy_metrics+1, cols=num_of_proxy+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_proxy):
        row.cells[item+1].text = proxynodes[item][20]

    for item in range(num_of_proxy_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(proxy_metrics[item])

        for proxynode in range(num_of_proxy):
            row.cells[proxynode+1].text = str(proxynodes[proxynode][item])


def vp():
    """
    Collect and collate the Qlik Sense Virtual Proxy information
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('Virtual Proxy details', 1)
    virtualproxy_metrics= ['description', 'prefix', 'authenticationModuleRedirectUri', 'sessionModuleBaseUri', 'loadBalancingModuleBaseUri', 'authenticationMethod', 'headerAuthenticationMode',
                           'headerAuthenticationHeaderName', 'headerAuthenticationStaticUserDirectory', 'headerAuthenticationDynamicUserDirectory', 'anonymousAccessMode', 'windowsAuthenticationEnabledDevicePattern',
                           'sessionCookieHeaderName', 'sessionCookieDomain', 'additionalResponseHeaders', 'sessionInactivityTimeout', 'extendedSecurityEnvironment', 'websocketCrossOriginWhiteList', 'defaultVirtualProxy',
                           'tags','samlMetadataIdP', 'samlHostUri', 'samlEntityId', 'samlAttributeUserId', 'samlAttributeUserDirectory', 'samlAttributeSigningAlgorithm', 'samlAttributeMap', 'magicLinkHostUri',
                           'magicLinkFriendlyName','jwtPublicKeyCertificate','jwtAttributeUserDirectory','jwtAttributeMap','jwtAttributeUserId']
    
    virtualproxynodes = get_qlik_sense.get_virtualproxy()
    num_of_virtualproxys = len(virtualproxynodes)
    num_of_virtualproxy_metrics = len(virtualproxy_metrics)
    table = document.add_table(rows=num_of_virtualproxy_metrics+1, cols=num_of_virtualproxys+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_virtualproxys):
        row.cells[item+1].text = virtualproxynodes[item][29]
    for item in range(num_of_virtualproxy_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(virtualproxy_metrics[item])
        for virtualproxynode in range(num_of_virtualproxys):
            row.cells[virtualproxynode+1].text = str(virtualproxynodes[virtualproxynode][item])

    document.add_page_break()

def vplb():
    """
    Collect and collate the Qlik Sense Virtual Proxy Load Balancer information
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 7772400
    section.page_height = 10058400
    document.add_heading('Virtual Proxy Load Balancing details', 1)
    virtualproxylbnodes = get_qlik_sense.get_vploadbalancers()
    virtualproxylb_metrics = ['node', 'load balance nodes']
    num_of_virtualproxyslb = len(virtualproxylbnodes)
    num_of_virtualproxylb_metrics = len(virtualproxylb_metrics)
    table = document.add_table(rows=num_of_virtualproxyslb+1, cols=2)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    for item in range(0, len(virtualproxylb_metrics)):
        row.cells[item].text = virtualproxylb_metrics[item]

    for item in range(num_of_virtualproxyslb):
        row = table.rows[item+1]
        row.cells[0].text = str(virtualproxylbnodes[item]['type'])
        row.cells[1].text = str(', '.join(virtualproxylbnodes[item]['items']))

    document.add_page_break()

def scheduler():
    """
    Collect and collate the Qlik Sense Scheduler information
    """
    document.add_heading('Scheduler details', 1)
    scheduler_metrics = ['customProperties', 'schedulerServiceType', 'maxConcurrentEngines', 'engineTimeout', 'tags', 'hostname',
                        'logVerbosityAuditActivity','logVerbosityAuditSecurity','logVerbosityService','logVerbosityApplication',
                        'logVerbosityAudit','logVerbosityPerformance','logVerbositySecurity','logVerbositySystem',
                        'logVerbosityTaskExecution']

    schedulernodes = get_qlik_sense.get_scheduler()
    num_of_schedulers = len(schedulernodes)
    num_of_scheduler_metrics = len(scheduler_metrics)
    table = document.add_table(rows=num_of_scheduler_metrics+1, cols=num_of_schedulers+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_schedulers):
        row.cells[item+1].text = schedulernodes[item][15]
    for item in range(num_of_scheduler_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(scheduler_metrics[item])
        for schedulernode in range(num_of_schedulers):
            row.cells[schedulernode+1].text = str(schedulernodes[schedulernode][item])
    document.add_page_break()

def apps():
    """
    Collect and collate the applications in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('Applications', level=1)
    apps = get_qlik_sense.get_apps()
    num_of_apps = len(apps)
    table = document.add_table(rows=num_of_apps+1, cols=7)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'App name'
    row.cells[1].text = 'App description'
    row.cells[2].text = 'Publish time'
    row.cells[3].text = 'Stream'
    row.cells[4].text = 'File size'
    row.cells[5].text = 'Owner userId'
    row.cells[6].text = 'Owner userName'
    for app in range(num_of_apps):
        row = table.rows[app+1]
        row.cells[0].text = str(apps[app][0])
        row.cells[1].text = str(apps[app][1])
        row.cells[2].text = str(apps[app][2])
        row.cells[3].text = str(apps[app][3])
        row.cells[4].text = str(apps[app][4])
        row.cells[5].text = str(apps[app][5])
        row.cells[6].text = str(apps[app][6])
    document.add_page_break()

def streams():
    """
    Collect and collate the streams in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('Streams', level=1)
    streams = get_qlik_sense.get_streams()
    num_of_streams = len(streams)
    table = document.add_table(rows=num_of_streams+1, cols=1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Stream name'
    for stream in range(num_of_streams):
        row = table.rows[stream+1]
        row.cells[0].text = str(streams[stream]['name'])
    document.add_page_break()

def data_connections():
    """
    Collect and collate the data connections in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('Data connections', level=1)
    connections = get_qlik_sense.get_dataconnections()
    num_of_connections = len(connections)
    table = document.add_table(rows=num_of_connections+1, cols=3)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'connection string'
    row.cells[2].text = 'type'
    for connection in range(num_of_connections):
        row = table.rows[connection+1]
        row.cells[0].text = str(connections[connection][0])
        row.cells[1].text = str(connections[connection][1])
        row.cells[2].text = str(connections[connection][2])
    document.add_page_break()

def user_directories():
    """
    Collect and collate the user directories in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('User Directories', level=1)
    userdirectories = get_qlik_sense.get_userdirectory()
    num_of_udc = len(userdirectories)
    table = document.add_table(rows=num_of_udc+1, cols=6)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'userDirectoryName'
    row.cells[2].text = 'configured'
    row.cells[3].text = 'operational'
    row.cells[4].text = 'type'
    row.cells[5].text = 'syncOnlyLoggedInUsers'
    for directory in range(num_of_udc):
        row = table.rows[directory+1]
        row.cells[0].text = str(userdirectories[directory][0])
        row.cells[1].text = str(userdirectories[directory][1])
        row.cells[2].text = str(userdirectories[directory][2])
        row.cells[3].text = str(userdirectories[directory][3])
        row.cells[4].text = str(userdirectories[directory][4])
        row.cells[5].text = str(userdirectories[directory][5])

    # document.add_page_break()

def users():
    """
    Collect and collate the users in the Qlik Sense Site (not currently written to document)
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('Users', level=1)
    users = get_qlik_sense.get_users()
    num_of_users = len(users)
    table = document.add_table(rows=num_of_users+1, cols=7)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'userId'
    row.cells[1].text = 'userDirectory'
    row.cells[2].text = 'name'
    row.cells[3].text = 'roles'
    row.cells[4].text = 'inactive'
    row.cells[5].text = 'removed externally'
    row.cells[6].text = 'blacklisted'
    for user in range(num_of_users):
        row = table.rows[user+1]
        row.cells[0].text = str(users[user][0])
        row.cells[1].text = str(users[user][1])
        row.cells[2].text = str(users[user][2])
        row.cells[3].text = str(users[user][3])
        row.cells[4].text = str(users[user][4])
        row.cells[5].text = str(users[user][5])
        row.cells[6].text = str(users[user][6])
    document.add_page_break()

def rules():
    """
    Collect and collate the system rules in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    
    document.add_heading('System Rules', level=1)
    document.add_heading('Rule Action Definition', level=2)
    paragraph = document.add_paragraph('The action on a rule is calculated by summing the selected permissions (e.g 3 = Create + Read.)')
    actions = [1, 2, 4, 8, 16, 32, 64, 128, 256]
    action_header = ['create', 'read', 'update', 'delete',
                     'export', 'publish', 'change owner',
                     'change role', 'export data']
    table = document.add_table(rows=2, cols=len(action_header))
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    for item in range(0, len(action_header)):
        row.cells[item].text = action_header[item]
    row = table.rows[1]
    row.cells[0].text = str(actions[0])
    row.cells[1].text = str(actions[1])
    row.cells[2].text = str(actions[2])
    row.cells[3].text = str(actions[3])
    row.cells[4].text = str(actions[4])
    row.cells[5].text = str(actions[5])
    row.cells[6].text = str(actions[6])
    row.cells[7].text = str(actions[7])
    row.cells[8].text = str(actions[8])
    paragraph = document.add_paragraph('')
    document.add_heading('Default rules', level=2)
    systemrules = get_qlik_sense.get_systemrules('Default')
    num_of_systemrules = len(systemrules)
    table = document.add_table(rows=num_of_systemrules+1, cols=5)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'rule'
    row.cells[2].text = 'resource filter'
    row.cells[3].text = 'actions'
    row.cells[4].text = 'disabled'
    for rule in range(num_of_systemrules):
        row = table.rows[rule+1]
        row.cells[0].text = str(systemrules[rule][0])
        row.cells[1].text = str(systemrules[rule][1])
        row.cells[2].text = str(systemrules[rule][2])
        row.cells[3].text = str(systemrules[rule][3])
        row.cells[4].text = str(systemrules[rule][4])

    document.add_page_break()
    document.add_heading('Custom rules', level=2)

    systemrules = get_qlik_sense.get_systemrules('Custom')
    num_of_systemrules = len(systemrules)
    table = document.add_table(rows=num_of_systemrules+1, cols=5)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'rule'
    row.cells[2].text = 'resource filter'
    row.cells[3].text = 'actions'
    row.cells[4].text = 'disabled'

    for rule in range(num_of_systemrules):
        row = table.rows[rule+1]
        row.cells[0].text = str(systemrules[rule][0])
        row.cells[1].text = str(systemrules[rule][1])
        row.cells[2].text = str(systemrules[rule][2])
        row.cells[3].text = str(systemrules[rule][3])
        row.cells[4].text = str(systemrules[rule][4])

    document.add_page_break()

    document.add_heading('Read only rules', level=2)

    systemrules = get_qlik_sense.get_systemrules('ReadOnly')
    num_of_systemrules = len(systemrules)
    table = document.add_table(rows=num_of_systemrules+1, cols=5)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'rule'
    row.cells[2].text = 'resource filter'
    row.cells[3].text = 'actions'
    row.cells[4].text = 'disabled'

    for rule in range(num_of_systemrules):
        row = table.rows[rule+1]
        row.cells[0].text = str(systemrules[rule][0])
        row.cells[1].text = str(systemrules[rule][1])
        row.cells[2].text = str(systemrules[rule][2])
        row.cells[3].text = str(systemrules[rule][3])
        row.cells[4].text = str(systemrules[rule][4])
    


    document.add_page_break()

def licenserules():
    """
    Collect and collate the license rules in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 10058400
    section.page_height = 7772400
    document.add_heading('License Rules', level=1)

    systemrules = get_qlik_sense.get_licenserules()
    num_of_systemrules = len(systemrules)
    table = document.add_table(rows=num_of_systemrules+1, cols=5)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'rule'
    row.cells[2].text = 'resource filter'
    row.cells[3].text = 'actions'
    row.cells[4].text = 'disabled'

    for rule in range(num_of_systemrules):
        row = table.rows[rule+1]
        row.cells[0].text = str(systemrules[rule][0])
        row.cells[1].text = str(systemrules[rule][1])
        row.cells[2].text = str(systemrules[rule][2])
        row.cells[3].text = str(systemrules[rule][3])
        row.cells[4].text = str(systemrules[rule][4])

    document.add_page_break()


def custom_props():
    """
    Collect and collate the custom properties in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 7772400
    section.page_height = 10058400
    document.add_heading('Custom Properties', level=1)

    customproperties = get_qlik_sense.get_customprop()
    num_of_customproperties = len(customproperties)
    table = document.add_table(rows=num_of_customproperties+1, cols=3)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    row.cells[1].text = 'choice values'
    row.cells[2].text = 'object types'

    for customproperty in range(num_of_customproperties):
        row = table.rows[customproperty+1]
        row.cells[0].text = str(customproperties[customproperty][0])
        row.cells[1].text = ', '.join(customproperties[customproperty][1])
        row.cells[2].text = ', '.join(customproperties[customproperty][2])
    document.add_page_break()

def tags():
    """
    Collect and collate the tags in the Qlik Sense Site
    """
    section = document.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = 7772400
    section.page_height = 10058400
    document.add_heading('Tags', level=1)
    tags = get_qlik_sense.get_tag()
    num_of_tags = len(tags)
    table = document.add_table(rows=num_of_tags+1, cols=1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'
    for tag in range(num_of_tags):
        row = table.rows[tag+1]
        row.cells[0].text = str(tags[tag])

def extensions():
    """
    Collect and collate the extensions in the Qlik Sense Site
    """
    document.add_page_break()
    document.add_heading('Extensions', level=1)
    extensions = get_qlik_sense.get_extensions()
    num_of_extensions = len(extensions)
    table = document.add_table(rows=num_of_extensions+1, cols=1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'name'

    for extension in range(num_of_extensions):
        row = table.rows[extension+1]
        row.cells[0].text = str(extensions[extension])

def qs_license():
    """
    Collect and collate the license in the Qlik Sense Site
    """
    paragraph = document.add_paragraph('')
    document.add_heading('License', level=1)
    lic_metric = ['lef', 'serial', 'name', 'organization', 'product', 'numberOfCores', 'isExpired', 'expiredReason', 'isBlacklisted', 'isInvalid']
    qs_lic = get_qlik_sense.get_license()
    num_of_metric = len(qs_lic)
    table = document.add_table(rows=num_of_metric+1, cols=2)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'details'

    for metric in range(len(qs_lic)):
        row = table.rows[metric+1]
        row.cells[0].text = str(lic_metric[metric])
        row.cells[1].text = str(qs_lic[metric][0])
    document.add_page_break()

def qs_servicecluster():
    """
    Collect and collate the service cluster in the Qlik Sense Site
    """
    paragraph = document.add_paragraph('')
    document.add_heading('Service Cluster', level=1)
    sc_metric = ['name', 'root folder', 'app folder', 'static content',
                 '32bit Connector', '64bit Connector', 'archived logs',
                 'database host', 'database port']
    qs_sc = get_qlik_sense.get_servicecluster()
    num_of_metric = len(qs_sc)
    table = document.add_table(rows=num_of_metric+1, cols=2)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'details'

    for metric in range(len(sc_metric)):
        row = table.rows[metric+1]
        row.cells[0].text = str(sc_metric[metric])
        row.cells[1].text = str(qs_sc[metric])
    document.add_page_break()

def node_config():
    """
    Collect and collate the node configuration in the Qlik Sense Site
    """
    paragraph = document.add_paragraph('')
    document.add_heading('Server node details', 1)
    node_metrics = ['host name','central','purpose', 'engine', 'proxy', 'printing', 'scheduler']
    nodes = get_qlik_sense.get_nodeconfig()
    num_of_nodes = len(nodes)
    num_of_node_metrics = len(node_metrics)
    table = document.add_table(rows=num_of_node_metrics+1, cols=num_of_nodes+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_nodes):
        row.cells[item+1].text = nodes[item][7]
    for item in range(num_of_node_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(node_metrics[item])

        for node in range(num_of_nodes):
            row.cells[node+1].text = str(nodes[node][item])
    document.add_page_break()

def repository_config():
    """
    Collect and collate the repository configuration in the Qlik Sense Site
    """
    paragraph = document.add_paragraph('')
    document.add_heading('Repository details', 1)
    node_metrics = ['logVerbosityAuditActivity',
                    'logVerbosityAuditSecurity',
                    'logVerbosityService',
                    'logVerbosityApplication',
                    'logVerbosityAudit',
                    'logVerbosityLicense',
                    'logVerbosityRuleAudit',
                    'logVerbosityManagementConsole',
                    'logVerbosityPerformance',
                    'logVerbositySecurity',
                    'logVerbositySynchronization',
                    'logVerbositySystem',
                    'logVerbosityserManagement']
    nodes = get_qlik_sense.get_logverbosity()
    num_of_nodes = len(nodes)
    num_of_node_metrics = len(node_metrics)
    table = document.add_table(rows=num_of_node_metrics+1, cols=num_of_nodes+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_nodes):
        row.cells[item+1].text = nodes[item][13]
    for item in range(num_of_node_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(node_metrics[item])
        for node in range(num_of_nodes):
            row.cells[node+1].text = str(nodes[node][item])
    document.add_page_break()

def printing():
    """
    Collect and collate the Qlik Sense Printing Service information
    """
    document.add_heading('Printing Service details', 1)

    printing_metrics = ['customproperties',
                         'workingSetSizeHiPct',
                         'logVerbosityAuditActivity',
                         'logVerbosityService',
                         'hostname',
                         'tags']

    printnodes = get_qlik_sense.get_printing()
    num_of_nodes = len(printnodes)
    num_of_print_metrics = len(printing_metrics)
    table = document.add_table(rows=num_of_print_metrics+1, cols=num_of_nodes+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_nodes):
        row.cells[item+1].text = printnodes[item][6]
    for item in range(num_of_print_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(printing_metrics[item])
        for printnode in range(num_of_nodes):
            row.cells[printnode+1].text = str(printnodes[printnode][item])

    document.add_page_break()

# def wmi():
#     """
#     Collect and collate the Windows Management Information from the servers that are involved in the Qlik Sense Site
#     """
#     try:
#         wminodes = server_details.auth.wmi_results()
#         document.add_heading('Windows Management Information details', 1)
#         wmi_metrics = ['operating system',
#                        'ram free',
#                        'total ram',
#                        'disk free',
#                        'disk total',
#                        'cpu name',
#                        'number of sockets',
#                        'number of logical cores']
#         num_of_wmis = len(wminodes)
#         num_of_wmi_metrics = len(wmi_metrics)
#         table = document.add_table(rows=num_of_wmi_metrics+1, cols=num_of_wmis+1)
#         table.style = 'Grid Table 1 Light Accent 1'
#         row = table.rows[0]
#         row.cells[0].text = 'Metric'
#         for item in range(0, num_of_wmis):
#             row.cells[item+1].text = wminodes[item][8]

#         for item in range(num_of_wmi_metrics):
#             row = table.rows[item+1]
#             row.cells[0].text = str(wmi_metrics[item])

#             for wminode in range(num_of_wmis):
#                 row.cells[wminode+1].text = str(wminodes[wminode][item])
#     except Exception:
#         print ('Failure to connect to wmi. > skipping wmi')

def userdirectorysettings():
    udcs = get_qlik_sense.get_udcdetails()
    if len(udcs) > 0:
        document.add_page_break()
        paragraph = document.add_paragraph('')
        document.add_heading('Details', 2)
        metrics = []
        udcs = get_qlik_sense.get_udcdetails()
        num_of_udcs = len(udcs)
        # get the length of each udc settings, this is used to create the number of table columns.
        for directory in udcs:
            metrics.append(len(udcs[directory]))

        table = document.add_table(rows=max(metrics)+1, cols=num_of_udcs+1)
        table.style = 'Grid Table 1 Light Accent 1'
        row = table.rows[0]
        row.cells[0].text = 'Metric'

        udc_names = []
        for item in udcs.keys():
            udc_names.append(item)
        for item in range(0, num_of_udcs):
            row.cells[item+1].text = udc_names[item]
        for item in range(max(metrics)):
            row = table.rows[item+1]
            row.cells[0].text = str('setting{0}'.format(item))

        udc_number = 1
        for directory in udc_names:
            for item in range(len(udcs[directory])):
                row = table.rows[item+1]
                row.cells[udc_number].text = str(udcs[directory]['setting{0}'.format(item)])
            udc_number += 1
        document.add_page_break()

def associatedproxy():
    paragraph = document.add_paragraph('')
    document.add_heading('Associated Virtual Proxy')
    assnproxy = get_qlik_sense.get_associatedproxies()
    num_of_proxy = len(assnproxy)
    table = document.add_table(rows=2, cols=num_of_proxy)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    colnum = 0
    for item in assnproxy:
        row.cells[colnum].text = item
        colnum += 1

    proxiescount = 0
    for item in assnproxy:
        row = table.rows[1]
        row.cells[proxiescount].text = str(', '.join(assnproxy[item]))    
        proxiescount+=1

def analytics():
    """
    Collect the advanced Analytics data from the site
    """
    document.add_heading('Analytic Connections', 1)

    ac_metrics = ['host',
                         'port',
                         'certificateFilePath',
                         'reconnectTimeout',
                         'requestTimeout']
    analyticConnections = get_qlik_sense.get_analyticconnections()
    num_of_conn = len(analyticConnections)
    num_of_analyticConnections_metrics = len(ac_metrics)
    table = document.add_table(rows=num_of_analyticConnections_metrics+1, cols=num_of_conn+1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Metric'
    for item in range(0, num_of_conn):
        row.cells[item+1].text = analyticConnections[item][5]
    for item in range(num_of_analyticConnections_metrics):
        row = table.rows[item+1]
        row.cells[0].text = str(ac_metrics[item])
        for conn in range(num_of_conn):
            row.cells[conn+1].text = str(analyticConnections[conn][item])
    document.add_page_break()

 
def odag():
    """
    Collect the odag data from the site
    """
    document.add_heading('OnDemand Application Generation (ODAG)', 1)

    odag_metrics = ['enabled',
                         'maxConcurrentRequests',
                         'logLevel']
    odag = get_qlik_sense.get_odag()
    num_of_metric = len(odag_metrics)
    table = document.add_table(rows=num_of_metric+1, cols=2)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'details'

    for metric in range(len(odag_metrics)):
        row = table.rows[metric+1]
        row.cells[0].text = str(odag_metrics[metric])
        row.cells[1].text = str(odag[0][metric])
    document.add_page_break()

def masterRoles():
    """
    Collect the master Roles from Qlik Sense
    """
    document.add_heading('Node Roles', 1)
    role= get_qlik_sense.masterRoles()
    num_of_nodes = len(role)
    table = document.add_table(rows=num_of_nodes+1, cols=2)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row = table.rows[0]
    row.cells[0].text = 'Host name'
    row.cells[1].text = 'Roles'
    for node in range(num_of_nodes):
        row = table.rows[node+1]
        row.cells[0].text = str(role[node][0])
        row.cells[1].text = str(role[node][1])

def test_connection():
    """
    Test the connection to the central node
    """
    try:
        connect()
    except:
        pass
        print ('Unable to connect.')
    else:
        main()

def savedoc():
    """
    Save output to word
    """
    document.save('QSDoc_{0}_{1}_{2}_{3}.docx'.format(args.server, year, month, day))

def main():
    """
    Main
    """
    print('Generating report..')
      
    dispatcher = [servernode,
                  node_config,
                  masterRoles,
                  qs_servicecluster,
                  qs_license,
                  repository_config,
                  engine,
                  proxy,
                  associatedproxy,
                  vp,
                  vplb,
                  scheduler,
                  printing,
                  analytics,
                  odag,
                  apps,
                  streams,
                  
                  data_connections,
                  user_directories,
                  userdirectorysettings,
                  licenserules,
                  rules,
                  custom_props,
                  tags,
                  extensions,
                  savedoc,
                  create_QSXls.main]
    if args.wmi:
        dispatcher = [servernode,
                      node_config,
                      masterRoles,
                      wmi,
                      qs_servicecluster,
                      qs_license,
                      repository_config,
                      engine,
                      proxy,
                      associatedproxy,
                      vp,
                      vplb,
                      scheduler,
                      printing,
                      analytics,
                      odag,
                      apps,
                      streams,

                      data_connections,
                      user_directories,
                      userdirectorysettings,
                      licenserules,
                      rules,
                      custom_props,
                      tags,
                      extensions,
                      savedoc,
                      create_QSXls.main]

    for item in progress.bar(range(len(dispatcher))):

        dispatcher[item]()

test_connection()


